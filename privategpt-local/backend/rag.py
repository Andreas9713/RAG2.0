from __future__ import annotations

import logging
from functools import lru_cache
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

from docx import Document as DocxDocument
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_community.llms import Ollama
from langchain_community.vectorstores import Chroma

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md"}
PERSIST_DIRECTORY = Path("./chroma")
DEFAULT_MODEL = "llama3.1"
FALLBACK_MODEL = "mistral"
OLLAMA_URL = "http://ollama:11434"


def _iter_source_files(paths: Sequence[str]) -> Iterable[Path]:
    for raw in paths:
        path = Path(raw).expanduser().resolve()
        if path.is_dir():
            for file_path in sorted(path.rglob("*")):
                if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                    yield file_path
        elif path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path
        else:
            logger.warning("Skipping unsupported path: %s", path)


def _load_docx(path: Path) -> List[Document]:
    doc = DocxDocument(str(path))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
    return [Document(page_content=text, metadata={"source": str(path)})]


def _load_document(path: Path) -> List[Document]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        loader = PyPDFLoader(str(path))
        return loader.load()
    if suffix == ".docx":
        return _load_docx(path)
    loader = TextLoader(str(path), encoding="utf-8")
    return loader.load()


def _split_documents(documents: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=120)
    return splitter.split_documents(documents)


@lru_cache(maxsize=1)
def _get_embeddings():
    return SentenceTransformerEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")


def _get_vectorstore() -> Chroma:
    PERSIST_DIRECTORY.mkdir(parents=True, exist_ok=True)
    return Chroma(persist_directory=str(PERSIST_DIRECTORY), embedding_function=_get_embeddings())


def ingest_paths(paths: Sequence[str]) -> int:
    documents: List[Document] = []
    for file_path in _iter_source_files(paths):
        try:
            documents.extend(_load_document(file_path))
        except Exception as exc:  # pragma: no cover - defensive
            logger.exception("Failed to load %s: %s", file_path, exc)
    if not documents:
        return 0

    chunks = _split_documents(documents)
    store = _get_vectorstore()
    store.add_documents(chunks)
    store.persist()
    return len(chunks)


def _build_prompt(context: str, question: str) -> str:
    template = (
        "You are a helpful AI assistant. Use the context to answer the question.\n"
        "If the answer is not in the context, reply that you don't know.\n"
        "Context:\n{context}\n\nQuestion: {question}\nAnswer:"
    )
    prompt = PromptTemplate.from_template(template)
    return prompt.format(context=context, question=question)


def _run_llm(prompt: str, model_name: str) -> str:
    llm = Ollama(model=model_name, base_url=OLLAMA_URL)
    return llm.invoke(prompt)


def _unique_sources(documents: Sequence[Document]) -> List[str]:
    seen: List[str] = []
    for doc in documents:
        source = doc.metadata.get("source")
        if source and source not in seen:
            seen.append(source)
    return seen


def ask(question: str, top_k: int = 5) -> Tuple[str, List[str]]:
    store = _get_vectorstore()
    retriever = store.as_retriever(search_kwargs={"k": top_k})
    docs = retriever.get_relevant_documents(question)
    if not docs:
        return "No information available in the knowledge base yet.", []

    context = "\n\n".join(doc.page_content for doc in docs)
    prompt = _build_prompt(context, question)

    try:
        answer = _run_llm(prompt, DEFAULT_MODEL)
    except Exception as exc:  # pragma: no cover - fallback path
        logger.warning("Primary model %s failed (%s), falling back to %s", DEFAULT_MODEL, exc, FALLBACK_MODEL)
        answer = _run_llm(prompt, FALLBACK_MODEL)

    sources = _unique_sources(docs)
    return answer.strip(), sources
